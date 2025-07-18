import os
import uuid
from pydantic import BaseModel
from llama_cpp import Llama
from docx import Document
import gradio as gr

# 加载 Qwen 模型
llm = Llama(
    model_path="models/Qwen3-0.6B-UD-Q4_K_XL.gguf",
    n_ctx=2048,
    n_threads=4,
    n_gpu_layers=0,
    verbose=False
)

# 定义数据模型类
class DocumentRequest(BaseModel):
    title: str
    context: str

def generate_text(prompt: str) -> str:
    """统一的文本生成函数"""
    output = llm(
        prompt ,
        max_tokens=1024,
        temperature=0.4,
        stop=["<|endoftext|>"]
    )
    return output['choices'][0]['text'].strip()

#前端：Gradio 页面
def gradio_analyze(file):
    file_id = str(uuid.uuid4())[:8]
    output_path = f"output/{file_id}_分析报告.docx"
    doc = Document(file.name)
    context = "\n".join([para.text for para in doc.paragraphs])
    prompt = f"""你是一名文档分析师，请分析以下内容：
1. 识别文档类型和核心主题
2. 用 150 字总结主要内容\n\n{context}\n\n分析报告："""
    output = llm( 
        prompt,
        max_tokens=1024,
        temperature=0.5,
        top_p=0.9,
        repeat_penalty=1.1, 
        stop=["<|endoftext|>"]
    )
    analysis = output['choices'][0]['text'].strip()
    result_doc = Document()
    result_doc.add_heading(f"文档分析报告 - {os.path.basename(file.name)}", 0)
    result_doc.add_paragraph(analysis)
    result_doc.save(output_path)
    
    return output_path

# 创建 Gradio 界面
gradio_app = gr.Interface(
    fn=gradio_analyze,
    inputs=gr.File(label="上传文档 (.docx)", file_types=[".docx"]),
    outputs=gr.File(label="下载分析报告"),
    title="📄 Autoch-alpha 文档智能分析系统--基于Qwen大模型",
    description="上传 Word 文档，AI 将自动分析内容并生成分析报告",
    allow_flagging="never"
)

# 启动服务
if __name__ == "__main__":
    import webbrowser
    webbrowser.open("http://localhost:7860") # 自动打开浏览器（好烦不搞还不能自己蹦出来）
    gradio_app.launch()
